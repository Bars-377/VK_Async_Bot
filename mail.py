# import threading

import smtplib
from email.message import EmailMessage
import time
import os
from pathlib import Path

host="172.18.11.104"
user="root"
password="enigma1418"
database="mdtomskbot"

def treatment_p(id, fio, contacts,type_service, service, date, category):

    if category == None:
        category = ''

    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    # Создайте новый документ
    doc = Document()

    # Функция для добавления текста по центру
    def add_centered_text(doc, text, bold=False, font_size=12):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Первая строка
    add_centered_text(doc, "ЗАЯВКА № ______", bold=True, font_size=14)

    date = str(date).split('-')

    # # Получение даты
    # date_str = date.strftime("от «%d»_%m_%Y г.")

    # Вторая строка
    add_centered_text(doc, f'от {date[2]}_{date[1]}_{date[0]} г.', font_size=12)

    # Третья строка
    add_centered_text(doc, "на оказание услуги по выездному обслуживанию заявителей", font_size=12)

    # Добавление таблицы
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    # Первая строка таблицы
    cell = table.cell(0, 0)
    cell.merge(table.cell(0, 1))
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = cell_paragraph.add_run("Сведения о заказчике")
    run.bold = True

    # Вторая строка таблицы
    table.cell(1, 0).text = "Фамилия Имя Отчество/ Наименование заказчика"

    # Вторая строка таблицы
    table.cell(1, 1).text = fio

    # Третья строка таблицы
    table.cell(2, 0).text = "Контактный телефон/ адрес проживания заказчика адрес электронной почты"

    # Третья строка таблицы
    table.cell(2, 1).text = contacts

    sender_email_address = "oprgp.toma@mfc.tomsk.ru"
    sender_password = "94NotYjc"  # Убедитесь, что это пароль приложения

    # Создайте сообщение электронной почты
    message = EmailMessage()
    message["Subject"] = f"Заявка на платную услугу от {fio}"
    message["From"] = sender_email_address
    if 'северск' in contacts.lower():
        message["To"] = 'zatoseversk@mfc.tomsk.ru'  # Адрес получателя
    else:
        message["To"] = 'msptosp@mfc.tomsk.ru'  # Адрес получателя

    # Установите текстовое содержимое письма
    message.set_content(f"Заявка на платную у слугу от {fio}")

    # Четвёртая строка таблицы
    table.cell(3, 0).text = "Наименование категории граждан, для которых организация выезда через МФЦ осуществляется бесплатно*"

    # Четвёртая строка таблицы
    table.cell(3, 1).text = category

    # Пятая строка таблицы
    table.cell(4, 0).text = "Вид услуги"
    if type_service == '1':
        table.cell(4, 1).text = "Выезд к заявителю с целью:\n⩗ приёма документов\n□ доставки документов"
    else:
        table.cell(4, 1).text = "Выезд к заявителю с целью:\n□ приёма документов\n⩗ доставки документов"

    # Шестая строка таблицы
    table.cell(5, 0).text = "Наименование услуги"

    # Шестая строка таблицы
    table.cell(5, 1).text = service

    # После таблицы
    doc.add_paragraph()
    add_centered_text(doc, "Заявку принял: __________(подпись) / __________(ФИО специалиста)", font_size=12)

    # Последняя строка справа
    paragraph = doc.add_paragraph("* - заполняется при выездном обслуживании льготных категорий граждан")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # # Сохраните документ
    # modified_attachment_path = "C:\\Users\\admin\\Desktop\\mail\\zayavka_" + str(id) + ".docx"

    # Получаем текущую директорию проекта
    # project_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = Path(__file__).resolve().parent

    # Строим путь к папке file внутри проекта
    # modified_attachment_path = os.path.join(project_dir, 'mail', f'zayavka_{str(id)}.docx')
    modified_attachment_path = project_dir / 'mail' / f'zayavka_{str(id)}.docx'

    doc.save(modified_attachment_path)

    print("Документ успешно создан и сохранён.")

    # Заполните эти переменные значениями для своего почтового сервера и учетной записи отправителя
    smtp_server = "smtp.mfc.tomsk.ru"
    # smtp_server = "smtp.yandex.ru"
    smtp_port = 587  # Порт для TLS

    # sender_email_address = "m.tosp@yandex.ru"
    # sender_password = "cwlecvijmxlpkvfo"  # Убедитесь, что это пароль приложения

    # # Прикрепите файл Excel к письму
    # attachment_path = "C:\\Users\\admin\\Desktop\\mail\\zayavka_" + str(id) + ".docx"

    # Получаем текущую директорию проекта
    # project_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = Path(__file__).resolve().parent

    # Строим путь к папке file внутри проекта
    # attachment_path = os.path.join(project_dir, 'mail', f'zayavka_{str(id)}.docx')
    attachment_path = project_dir / 'mail' / f'zayavka_{str(id)}.docx'

    # with open(attachment_path, "rb") as attachment:
    #     message.add_attachment(attachment.read(), maintype="application", subtype="vnd.openxmlformats-officedocument.wordprocessingml.document", filename=os.path.basename(attachment_path))

    with attachment_path.open("rb") as attachment:
        message.add_attachment(
            attachment.read(),
            maintype="application",
            subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=attachment_path.name
        )

    # Функция для отправки письма
    def send_email():
        condition = False
        try:
            # Подключитесь к SMTP-серверу
            with smtplib.SMTP(smtp_server, smtp_port) as server:
                server.starttls()  # Инициализация TLS
                server.login(sender_email_address, sender_password)
                server.send_message(message)  # Отправьте сообщение
            print("Письмо успешно отправлено.")

            # Удалите документ
            # os.remove(modified_attachment_path)
            modified_attachment_path.unlink()

            condition = True
        except smtplib.SMTPAuthenticationError as e:
            print(f"Ошибка аутентификации mail: {e}")
        except Exception as e:
            print(f"Произошла ошибка mail: {e}")

        return condition

    # Попробуйте отправить письмо несколько раз с интервалами
    attempts = 3
    for attempt in range(attempts):
        print(f"Попытка {attempt + 1} из {attempts}")

        condition = send_email()
        if condition:
            break
        elif attempt < attempts - 1:
            print("Ожидание перед повторной попыткой...")
            time.sleep(10)  # Ожидание 10 секунд перед следующей попыткой

    return

import mysql.connector

def process_mail():

    # Создание одиночного соединения
    dbconfig = {
        'host': host,
        'user': user,
        'password': password,
        'database': database,
    }

    connection = mysql.connector.connect(**dbconfig)

    # Выполнение запросов и обновление данных
    try:
        cursor = connection.cursor()

        while True:
            # global exit_event
            update_query = "SELECT * FROM application;"
            cursor.execute(update_query)
            data = cursor.fetchall()

            for row in data:
                treatment_p(row[0], row[1], row[2], row[3], row[4], row[5], row[6])
                update_query = f"DELETE FROM application WHERE id = {row[0]};"
                cursor.execute(update_query)

            connection.commit()  # Фиксация транзакции

    except Exception as e:
        # Вывод подробной информации об ошибке
        print(f"Поймано исключение: {type(e).__name__}")
        print(f"Сообщение об ошибке: {str(e)}")
        import traceback
        print("Трассировка стека (stack trace):")
        traceback.print_exc()

    finally:
        if 'cursor' in locals() or 'cursor' in globals():
            cursor.close()  # Важно закрыть курсор
        if 'connection' in locals() or 'connection' in globals():
            connection.close()  # Важно закрыть соединение после его использования

def process_file():
    import requests

    import configparser
    # Чтение конфигурации
    config = configparser.ConfigParser()
    config.read("config.ini")

    # Ваш токен бота
    TOKEN = config["TELEGRAM"]["token_one"]

    def search_user_id():
        import mysql.connector
        try:
            mydb = mysql.connector.connect(
                host=host,
                user=user,
                password=password,
                database=database,
                connect_timeout=2
            )
            mycursor = mydb.cursor()

            # Search in the database using ani and file_id
            query = f"SELECT id_tb, phone FROM file_id_tb WHERE marker = 'yes';"
            mycursor.execute(query)
            result = mycursor.fetchone()

            mydb.close()

            if not result:
                return None, None

            if result[0] == '':
                return None, None
            elif result[1] == '':
                return None, None
            elif result[0] == '' and result[1] == '':
                return None, None

            # Возвращаем результат
            return result[0], result[1]
        except Exception as e:
            print('search_user_id:', e)
            return None, None

    def search_files(user_id):
        import mysql.connector
        try:
            mydb = mysql.connector.connect(
                host=host,
                user=user,
                password=password,
                database=database,
                connect_timeout=2
            )
            mycursor = mydb.cursor()

            # Search in the database using ani and file_id
            query = f"SELECT file, file_format, message FROM file_id_tb WHERE id_tb = '{user_id}' AND message IN ('salesman_passport', 'buyer_passport', 'proxy_passport', 'EGRN', 'right');"

            mycursor.execute(query)
            result = mycursor.fetchall()

            mydb.close()

            return result if result else None

        except Exception as e:
            print('search_files:', e)
            return None

    def search_message(user_id):
        import mysql.connector
        try:
            mydb = mysql.connector.connect(
                host=host,
                user=user,
                password=password,
                database=database,
                connect_timeout=2
            )
            mycursor = mydb.cursor()

            # Search in the database using ani and file_id
            query = f"SELECT message FROM file_id_tb WHERE id_tb = '{user_id}' AND message LIKE '%Цена%';"
            mycursor.execute(query)
            result = mycursor.fetchall()

            mydb.close()

            return result[0][0] if result else None

        except Exception as e:
            print('search_message:', e)
            return None

    def delete_files(user_id):
        import mysql.connector
        try:
            mydb = mysql.connector.connect(
                host=host,
                user=user,
                password=password,
                database=database,
                connect_timeout=2
            )
            mycursor = mydb.cursor()

            # Search in the database using ani and file_id
            query = f"DELETE FROM file_id_tb WHERE id_tb = '{user_id}';"
            mycursor.execute(query)
            mydb.commit()

            mydb.close()

            return

        except Exception as e:
            print('delete_files:', e)
            return

    def send_files():
        # while True:
        try:
            user_id, phone = search_user_id()

            if user_id and phone:
                files = search_files(user_id)

                # # Создаем папку "321", если она не существует
                # save_directory = "C:\\Users\\admin\\Desktop\\file"

                # Получаем текущую директорию проекта
                # project_dir = os.path.dirname(os.path.abspath(__file__))
                project_dir = Path(__file__).resolve().parent

                # Строим путь к папке file внутри проекта
                # save_directory = os.path.join(project_dir, 'file')
                save_directory = project_dir / 'file'

                if files:
                    for file in files:

                        if file[0] and file[1] and file[2]:

                            file_name = f'{file[2]}_{user_id}.{file[1]}'

                            # elif "photo" in message:
                            #     # Выбираем последнее (самое большое) фото
                            #     largest_photo = message["photo"][-1]
                            #     file_id = largest_photo["file_id"]
                            #     file_name = f"{file_id}.jpg"
                            #     print(f"file_id фотографии: {file_id}")

                            # elif "video" in message:
                            #     file_id = message["video"]["file_id"]
                            #     file_name = f"{file_id}.mp4"
                            #     print(f"file_id видео: {file_id}")

                            # Получаем путь к файлу
                            file_info_url = f"https://api.telegram.org/bot{TOKEN}/getFile?file_id={file[0]}"
                            file_info_response = requests.get(file_info_url)
                            file_info = file_info_response.json()

                            if file_info.get("ok"):
                                file_path = file_info['result']['file_path']

                                # Формируем URL для загрузки файла
                                file_url = f"https://api.telegram.org/file/bot{TOKEN}/{file_path}"

                                # Загружаем файл
                                file_response = requests.get(file_url)

                                # Создайте директорию, если она не существует
                                # os.makedirs(save_directory, exist_ok=True)
                                save_directory.mkdir(parents=True, exist_ok=True)

                                # Путь к сохранённому файлу
                                # save_path = os.path.join(save_directory, file_name)
                                save_path = save_directory / file_name

                                # Сохраните содержимое файла
                                # with open(save_path, "wb") as f:
                                #     f.write(file_response.content)
                                with save_path.open("wb") as f:
                                    f.write(file_response.content)

                                print(f"Файл сохранен по пути: {save_path}")

                            else:
                                print("Ошибка при получении пути к файлу:", file_info)

                message = search_message(user_id)
                # Разбиваем строку по символу "_"
                if message:
                    split_text = message.split('_')
                    # Соединяем элементы попарно с ": "
                    result = [f"{split_text[i]} {split_text[i+1]}" for i in range(0, len(split_text)-1, 2)]
                    # print(result)
                    result.append(phone)
                    result = ', '.join(result)

                    # file_path = os.path.join(save_directory, f'info_{user_id}.txt')  # Путь к файлу
                    file_path = save_directory / f'info_{user_id}.txt' # Путь к файлу

                    # Записываем file_path в файл
                    with open(file_path, 'w') as file:
                        file.write(result)

                delete_files(user_id)

            import smtplib
            import time
            from email.mime.base import MIMEBase
            from email import encoders

            # # Путь к папке, где находятся файлы
            # folder_path = "C:\\Users\\admin\\Desktop\\file"

            # Получаем текущую директорию проекта
            # project_dir = os.path.dirname(os.path.abspath(__file__))
            project_dir = Path(__file__).resolve().parent

            # Строим путь к папке file внутри проекта
            # folder_path = os.path.join(project_dir, 'file')
            folder_path = project_dir / 'file'

            # Функция для поиска первого файла, содержащего "info" в имени
            # def find_file_with_info(folder):
            #     for filename in os.listdir(folder):
            #         if f"info" in filename:
            #             return os.path.join(folder, filename), str(filename).split('_')[1].split('.')[0]
            #     return None, None
            def find_file_with_info(folder):
                for file in folder.iterdir():
                    if file.is_file() and "info" in file.name:
                        return file, file.name.split('_')[1].split('.')[0]
                return None, None

            # Найдите файл с "info" в имени
            file_with_info, user_id = find_file_with_info(folder_path)

            if not file_with_info:
                # print("Файл, содержащий 'info', не найден.")
                return

            # Функция для извлечения текста из файла
            def extract_text_from_file(filepath):
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        return f.read()
                except:
                    with open(filepath, 'r', encoding='windows-1251') as f:
                        return f.read()

            # Извлекаем текст из файла
            file_text = extract_text_from_file(file_with_info)

            # Функция для получения первой части имени файла до символа "_"
            def extract_search_string_from_filename(found_files):
                search_strings = []
                for file in found_files:
                    search_strings.append(str(file).split("\\")[-1].split('.')[0] if "_" in str(file) else None)
                return search_strings if search_strings else None

            # Функция для поиска первого файла
            # def get_files_in_folder(folder,user_id):
            #     found_files = []
            #     files = os.listdir(folder)
            #     for filename in files:
            #         if user_id in filename:
            #             found_files.append(os.path.join(folder, filename))
            #     return found_files if found_files else None
            def get_files_in_folder(folder, user_id):
                found_files = []
                for file in folder.iterdir():
                    if file.is_file() and user_id in file.name:
                        found_files.append(file)
                return found_files if found_files else None

            # Найдите первый файл в папке
            found_files = get_files_in_folder(folder_path, user_id)

            if not found_files:
                # print("В папке нет файлов.")
                return

            # Извлекаем строку для поиска из имени файла
            search_strings = extract_search_string_from_filename(found_files)

            if not search_strings:
                # print(f"Не удалось извлечь строку из имени файла: {found_files}")
                return

            """ТУТ КОСЯК"""
            # def find_file_in_folder(folder, search_strings):
            #     files_found = []
            #     files = os.listdir(folder)  # Список файлов в папке

            #     for filename in files:
            #         # Преобразуем имя файла в нижний регистр и проверяем условия
            #         if 'info' not in filename and str(filename).split('.')[0] in search_strings:
            #             files_found.append(os.path.join(folder, filename))

            #     return files_found if files_found else None

            def find_file_in_folder(folder, search_strings):
                files_found = []
                search_strings_transformed = [Path(path).stem for path in search_strings]
                for file in folder.iterdir():  # Список файлов в папке
                    if file.is_file():  # Проверка, чтобы работать только с файлами
                        # Преобразуем имя файла в нижний регистр и проверяем условия
                        if 'info' not in file.name and file.stem in search_strings_transformed:
                            files_found.append(file)

                return files_found if files_found else None

            # Найдите файл, содержащий search_string
            modified_attachment_path = find_file_in_folder(folder_path, search_strings)

            # if not modified_attachment_path:
                # print(f"Файл, содержащий '{search_strings}' в имени, не найден.")
                # return

            # Настройки почтового сервера
            smtp_server = "smtp.mfc.tomsk.ru"
            smtp_port = 587  # Порт для TLS
            sender_email_address = "oprgp.toma@mfc.tomsk.ru"
            sender_password = "94NotYjc"  # Убедитесь, что это пароль приложения

            from email.mime.multipart import MIMEMultipart

            # Создайте сообщение электронной почты
            # message = EmailMessage()
            message = MIMEMultipart()
            message["Subject"] = f"Документы на составление договора от {file_text.split(',')[-2]}"
            message["From"] = sender_email_address
            # message["To"] = 'msptosp@mfc.tomsk.ru'  # Адрес получателя

            filials_id = {
                        '533': 'kirovskiy@mfc.tomsk.ru',
                        '461': 'leninskiy@mfc.tomsk.ru',
                        '641': 'oktyabrskiy@mfc.tomsk.ru',
                        '689': 'sovetskiy@mfc.tomsk.ru',
                        '431': 'zatoseversk@mfc.tomsk.ru',
                        # '425': 'gulyaeva@mfc.tomsk.ru',
                        # '665': 'msptosp@mfc.tomsk.ru',
                        # '377': 'msptosp@mfc.tomsk.ru',
                        # '557': 'msptosp@mfc.tomsk.ru',
                        # '389': 'msptosp@mfc.tomsk.ru',
                        # '449': 'msptosp@mfc.tomsk.ru',
                        # '677': 'msptosp@mfc.tomsk.ru',
                        # '605': 'msptosp@mfc.tomsk.ru',
                        '443': 'msptosp@mfc.tomsk.ru',
                        # '719': 'msptosp@mfc.tomsk.ru',
                        # '659': 'msptosp@mfc.tomsk.ru',
                        '479': 'msptosp@mfc.tomsk.ru',
                        # '713': 'msptosp@mfc.tomsk.ru',
                        '731': 'msptosp@mfc.tomsk.ru',
                        '551': 'msptosp@mfc.tomsk.ru',
                        # '407': 'msptosp@mfc.tomsk.ru',
                        '725': 'msptosp@mfc.tomsk.ru',
                        # '467': 'msptosp@mfc.tomsk.ru',
                        # '623': 'msptosp@mfc.tomsk.ru',
                        # '347': 'msptosp@mfc.tomsk.ru',
                        # '545': 'msptosp@mfc.tomsk.ru',
                        # '635': 'msptosp@mfc.tomsk.ru',
                        # '455': 'msptosp@mfc.tomsk.ru',
                        '395': 'msptosp@mfc.tomsk.ru',
                        '491': 'msptosp@mfc.tomsk.ru',
                        '371': 'msptosp@mfc.tomsk.ru',
                        # '611': 'msptosp@mfc.tomsk.ru',
                        # '593': 'msptosp@mfc.tomsk.ru',
                        # '503': 'msptosp@mfc.tomsk.ru',
                        # '443': 'msptosp@mfc.tomsk.ru',
                        '401': 'msptosp@mfc.tomsk.ru',
                        # '629': 'msptosp@mfc.tomsk.ru',
                        '539': 'msptosp@mfc.tomsk.ru',
                        # '2894026': 'msptosp@mfc.tomsk.ru',
                        # '719': 'msptosp@mfc.tomsk.ru',
                        # '515': 'msptosp@mfc.tomsk.ru',
                        '575': 'msptosp@mfc.tomsk.ru',
                        '437': 'msptosp@mfc.tomsk.ru',
                        # '647': 'msptosp@mfc.tomsk.ru',
                        # '707': 'msptosp@mfc.tomsk.ru',
                        # '479': 'msptosp@mfc.tomsk.ru',
                        '383': 'msptosp@mfc.tomsk.ru',
                        '329': 'msptosp@mfc.tomsk.ru',
                        '419': 'msptosp@mfc.tomsk.ru',
                        # '563': 'msptosp@mfc.tomsk.ru',
                        '599': 'msptosp@mfc.tomsk.ru',
                        # '335': 'msptosp@mfc.tomsk.ru',
                        # '617': 'msptosp@mfc.tomsk.ru',
                        # '497': 'msptosp@mfc.tomsk.ru',
                        # '407': 'msptosp@mfc.tomsk.ru',
                        # '413': 'msptosp@mfc.tomsk.ru',
                        '527': 'msptosp@mfc.tomsk.ru',
                        # '467': 'msptosp@mfc.tomsk.ru',
                        '521': 'msptosp@mfc.tomsk.ru',
                        # '653': 'msptosp@mfc.tomsk.ru',
                        # '671': 'msptosp@mfc.tomsk.ru',
                        # '473': 'msptosp@mfc.tomsk.ru',
                        '623': 'msptosp@mfc.tomsk.ru',
                        # '569': 'msptosp@mfc.tomsk.ru',
                        # '485': 'msptosp@mfc.tomsk.ru',
                        # '359': 'msptosp@mfc.tomsk.ru',
                        # '347': 'msptosp@mfc.tomsk.ru',
                        # '371': 'msptosp@mfc.tomsk.ru',
                        # '342595': 'msptosp@mfc.tomsk.ru',
                        '683': 'msptosp@mfc.tomsk.ru',
                        # '449': 'msptosp@mfc.tomsk.ru',
                        # '432256': 'msptosp@mfc.tomsk.ru',
                        # '365': 'msptosp@mfc.tomsk.ru',
                        # '695': 'msptosp@mfc.tomsk.ru',
                        # '701': 'msptosp@mfc.tomsk.ru',
                        '509': 'msptosp@mfc.tomsk.ru',
                        # '432194': 'msptosp@mfc.tomsk.ru',
                        # '432212': 'msptosp@mfc.tomsk.ru',
                        # '432239': 'msptosp@mfc.tomsk.ru',
                        'oprgp.toma': 'oprgp.toma@mfc.tomsk.ru'
                    }

            # message["To"] = 'gulyaeva@mfc.tomsk.ru'  # Адрес получателя
            try:
                message["To"] = filials_id[file_text.split(',')[-1].strip()]
            except:
                if modified_attachment_path:
                    # Удалите документ после отправки
                    for file in modified_attachment_path:
                        # os.remove(file)
                        file.unlink()
                # os.remove(file_with_info)
                file_with_info.unlink()
                return

            if message["To"] == 'oprgp.toma@mfc.tomsk.ru':
                message["Subject"] = f"Техподдержка от {file_text.split(',')[-2]}"

            if file_text.split(',')[-2] == ' 999':
                message["Subject"] = f"Чек"

            if 'Услуги СВО' in file_text:
                message["Subject"] = f"Услуги СВО от {file_text.split(',')[-2]}"

            file_text = file_text.split(',')

            if message["To"] == 'oprgp.toma@mfc.tomsk.ru':
                # Фильтрация массива: удаляем только слово "Цена", оставляя "Привет"
                file_text = [item.replace("Цена", "").strip() for item in file_text]

            if any('Услуги СВО' in item for item in file_text):
                # Фильтрация массива: удаляем только слово "Цена", оставляя "Привет"
                file_text = [item.replace("Цена", "").strip() for item in file_text]

            if any(' 999' in item for item in file_text):
                # Фильтрация массива: удаляем только слово "Цена", оставляя "Привет"
                file_text = [item.replace("Цена", "").strip() for item in file_text]
                file_text = [item.replace("999", "").strip() for item in file_text]
                file_text = file_text[:-1] # Удаляем последний элемент

            file_text = file_text[:-1] # Удаляем последний элемент

            if message["To"] == 'oprgp.toma@mfc.tomsk.ru':
                file_text = file_text[:-1] # Удаляем последний элемент

            if any('Услуги СВО' in item for item in file_text):
                file_text = file_text[:-1] # Удаляем последний элемент

            file_text = ','.join(file_text)

            # Установите текстовое содержимое письма
            # message.set_content(file_text)
            from email.mime.text import MIMEText
            message.attach(MIMEText(file_text, 'plain'))

            # Определение MIME-типов для произвольных файлов
            def guess_mime_type(filepath):
                import mimetypes
                mimetype, _ = mimetypes.guess_type(filepath)
                return mimetype or 'application/octet-stream'

            # Функция для прикрепления файла
            def attach_file(message, file_path):
                for file in file_path:
                    print(file)
                    mime_type = guess_mime_type(file)
                    maintype, subtype = mime_type.split('/', 1)
                    # with open(file, 'rb') as attachment:
                    #     part = MIMEBase(maintype, subtype)
                    #     part.set_payload(attachment.read())
                    #     encoders.encode_base64(part)
                    #     part.add_header('Content-Disposition', 'attachment', filename=os.path.basename(file))
                    #     message.attach(part)
                    with file.open("rb") as attachment:
                        part = MIMEBase(maintype, subtype)
                        part.set_payload(attachment.read())
                        encoders.encode_base64(part)
                        part.add_header('Content-Disposition', 'attachment', filename=file.name)
                        message.attach(part)

            if modified_attachment_path:
                attach_file(message, modified_attachment_path)

            # Функция для отправки письма
            def send_email():
                condition = False
                try:
                    # Подключитесь к SMTP-серверу
                    with smtplib.SMTP(smtp_server, smtp_port) as server:
                        server.starttls()  # Инициализация TLS
                        server.login(sender_email_address, sender_password)
                        server.send_message(message)  # Отправьте сообщение
                    print(f"Письмо успешно отправлено на {message["To"]}.")

                    if modified_attachment_path:
                        # Удалите документ после отправки
                        for file in modified_attachment_path:
                            # os.remove(file)
                            file.unlink()
                    # os.remove(file_with_info)
                    file_with_info.unlink()

                    condition = True
                except smtplib.SMTPAuthenticationError as e:
                    print(f"Ошибка аутентификации mail: {e}")
                except Exception as e:
                    print(f"Произошла ошибка mail: {e}")

                return condition

            # Попробуйте отправить письмо несколько раз с интервалами
            attempts = 3
            for attempt in range(attempts):
                print(f"Попытка {attempt + 1} из {attempts}")

                condition = send_email()
                if condition:
                    break
                elif attempt < attempts - 1:
                    print("Ожидание перед повторной попыткой...")
                    time.sleep(10)  # Ожидание 10 секунд перед следующей попыткой

            return

        except Exception as e:
            # Вывод подробной информации об ошибке
            print(f"Поймано исключение: {type(e).__name__}")
            print(f"Сообщение об ошибке: {str(e)}")
            import traceback
            print("Трассировка стека (stack trace):")
            traceback.print_exc()

    send_files()

# if __name__ == "__main__":

#     # создание и запуск потоков для каждого процесса
#     t1 = threading.Thread(target=process_mail)

#     while True:
#         if not t1.is_alive():
#             t1 = threading.Thread(target=process_mail)
#             t1.start()
